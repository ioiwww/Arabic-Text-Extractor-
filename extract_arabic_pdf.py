import os
import re
import pytesseract
from pdfminer.high_level import extract_text
from pdf2image import convert_from_path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import arabic_reshaper
# ألغينا استيراد bidi لأنها هي سبب القلب الزائد

INPUT_FOLDER = "pdfs"
OUTPUT_FOLDER = "output"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# نمط اكتشاف اللغة العربية
arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')

def fix_arabic_text(text):
    if not text.strip():
        return ""
    # "لحم" الحروف ببعضها لتبدو كلمات صحيحة
    # سنلغي استخدام get_display هنا لترك المتصفح أو الوورد يحدد الاتجاه
    reshaped_text = arabic_reshaper.reshape(text)
    return reshaped_text

def extract_arabic_lines(text):
    lines = text.split("\n")
    arabic_lines = []
    for line in lines:
        if arabic_pattern.search(line):
            # نصلح شكل الحروف فقط
            arabic_lines.append(fix_arabic_text(line.strip()))
    return arabic_lines

def extract_pdf_text(pdf):
    try:
        return extract_text(pdf)
    except:
        return ""

def extract_scanned_text(pdf):
    # تحويل الـ PDF لصور ومعالجتها بـ OCR
    try:
        pages = convert_from_path(pdf)
        text = ""
        for page in pages:
            text += pytesseract.image_to_string(page, lang="ara") + "\n"
        return text
    except Exception as e:
        print(f"Error in OCR: {e}")
        return ""

for file in os.listdir(INPUT_FOLDER):
    if not file.endswith(".pdf"): continue
    
    pdf_path = os.path.join(INPUT_FOLDER, file)
    print(f"Processing: {file}")

    # محاولة استخراج النص الرقمي
    text = extract_pdf_text(pdf_path)
    
    # إذا النص قليل أو غير مفهوم نستخدم الـ OCR
    if len(text.strip()) < 200:
        print("Using OCR...")
        text = extract_scanned_text(pdf_path)

    arabic_lines = extract_arabic_lines(text)
    base = os.path.splitext(file)[0]
    
    # 1. حفظ ملف TXT
    txt_file = os.path.join(OUTPUT_FOLDER, base + "_arabic.txt")
    with open(txt_file, "w", encoding="utf-8") as f:
        # هنا قد نحتاج لقلب الاتجاه للعرض في Notepad فقط
        f.write("\n".join(arabic_lines))

    # 2. حفظ ملف DOCX (وهو الأهم للتنسيق)
    docx_file = os.path.join(OUTPUT_FOLDER, base + "_arabic.docx")
    doc = Document()
    
    for line in arabic_lines:
        p = doc.add_paragraph()
        # إضافة النص وتحديد المحاذاة لليمين
        run = p.add_run(line)
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        
    doc.save(docx_file)
    print(f"Saved: {base}")

print("Done!")
